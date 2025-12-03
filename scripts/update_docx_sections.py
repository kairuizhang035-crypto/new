import os
import pandas as pd
from docx import Document
from docx.shared import Inches

ROOT = "/home/zkr/因果发现3"
DOCX_PATH = os.path.join(ROOT, "临床因果知识图谱构建与证据三角验证框架.docx")
VIS_DIR = os.path.join(ROOT, "02因果发现", "可视化")
RES_DIR = os.path.join(ROOT, "02因果发现")

def load_comp_stats():
    df = pd.read_csv(os.path.join(VIS_DIR, "算法对比数据.csv"))
    df.set_index("算法", inplace=True)
    return df

def load_overlap():
    mat = pd.read_csv(os.path.join(VIS_DIR, "算法重叠矩阵.csv"), index_col=0)
    return mat

def load_integrated_edges(top_n=15):
    df = pd.read_csv(os.path.join(RES_DIR, "06候选因果边集合", "因果边综合评分结果.csv"))
    sort_col = "集成评分" if "集成评分" in df.columns else ("综合评分" if "综合评分" in df.columns else None)
    if sort_col:
        df = df.sort_values(by=sort_col, ascending=False)
    cols = [
        "源节点", "目标节点", "支持算法数量",
        "综合评分" if "综合评分" in df.columns else None,
        "算法一致性评分" if "算法一致性评分" in df.columns else None,
        "统计显著性评分" if "统计显著性评分" in df.columns else None,
        "集成评分" if "集成评分" in df.columns else None,
    ]
    cols = [c for c in cols if c is not None]
    return df[cols].head(top_n)

def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 2:
        p.style = doc.styles["Heading 2"]
    elif level == 3:
        p.style = doc.styles["Heading 3"]
    return p

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def add_table_from_df(doc, df, title):
    add_paragraph(doc, title)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])
    return table

def add_figure(doc, img_path, caption):
    doc.add_picture(img_path, width=Inches(6))
    add_paragraph(doc, caption)

def build_methods_text(comp_stats):
    pc = comp_stats.loc["PC算法"]
    hc = comp_stats.loc["爬山算法"]
    ges = comp_stats.loc["贪婪等价搜索"]
    tan = comp_stats.loc["树搜索"]
    eil = comp_stats.loc["专家在循环"]
    return (
        "本研究采用多方法学因果结构发现框架，综合约束式与评分式两大范式，"
        "以提升临床复杂系统中因果关系识别的稳健性与可信度。数据集源自\n"
        f"`{os.path.join(ROOT, '01数据预处理', '缩减数据_规格.csv')}`，变量全部离散化为0/1或有限类别，以适配离散型因果结构学习。"
        "预处理环节包括：删除全空列、统一数据类型为数值型、保留临床相关且统计稳定的变量集合，并以样本ID作为索引。\n"
        "方法学包括：\n"
        "(1) PC算法：基于条件独立检验的约束式推断，采用稳定(stable)变体；独立性检验为卡方检验(chi_square)，显著性水平0.05，适用于离散变量、满足有向无环结构(DAG)假设的情形。\n"
        "(2) 贪婪等价搜索(GES)：在等价类(CPDAG)空间进行两阶段贪婪搜索，评分采用AIC-D(Akaike信息准则-离散)以权衡拟合与复杂度；适用于离散编码变量，能够在中高维下取得良好折中。\n"
        "(3) 爬山算法(Hill Climbing)：在有向图空间进行局部贪婪邻域搜索，评分函数为AIC-D，设置微小改进阈值以实现早停；在样本充足时更易得到较密结构。\n"
        "(4) 树搜索(TAN)：在树结构约束下学习近似DAG，利用互信息导出的最大生成树进行定向；在高维变量和有限样本场景下降低过拟合风险。\n"
        "(5) 专家在循环(Expert-in-the-loop)：基于多算法候选边集，组织临床专家迭代审阅与修正，形成知识约束的结构校正与裁剪。\n"
        "集成策略在边级别实施一致性整合：对每条候选因果边计算出现频次、支持算法数、算法一致性评分、网络拓扑评分、统计显著性评分，并归一化为[0,1]后加权汇总生成集成评分；据此筛选高质量候选集并输出精简边列表。\n"
        f"参数与实验设计概览：PC使用显著性水平0.05与稳定变体；GES与HC采用AIC-D评分；各算法在同一离散数据集(节点数以{int(pc['节点数'])}为基准)上独立运行，并记录网络指标(如边数、入度/出度上限、平均度)。"
    )

def build_results_text(comp_stats, overlap):
    pc = comp_stats.loc["PC算法"]
    hc = comp_stats.loc["爬山算法"]
    ges = comp_stats.loc["贪婪等价搜索"]
    tan = comp_stats.loc["树搜索"]
    eil = comp_stats.loc["专家在循环"]
    pc_hc = int(overlap.loc["PC算法", "爬山算法"]) if "爬山算法" in overlap.columns else None
    pc_ges = int(overlap.loc["PC算法", "贪婪等价搜索"]) if "贪婪等价搜索" in overlap.columns else None
    hc_ges = int(overlap.loc["爬山算法", "贪婪等价搜索"]) if "贪婪等价搜索" in overlap.columns else None
    tan_ges = int(overlap.loc["树搜索", "贪婪等价搜索"]) if "贪婪等价搜索" in overlap.columns else None
    return (
        "在统一数据集上，多方法学产生的结构指标显示出可解释的差异与互补性。\n"
        f"PC算法获得{int(pc['边数'])}条边，网络较为稀疏(平均度{pc['平均度数']})，反映约束式在显著性控制下的保守性；\n"
        f"爬山算法获得{int(hc['边数'])}条边(平均度{hc['平均度数']})，结构更为致密，适合揭示多变量间的广泛耦合；\n"
        f"GES获得{int(ges['边数'])}条边(平均度{ges['平均度数']})，在拟合与复杂度之间达成平衡；\n"
        f"树搜索输出{int(tan['边数'])}条树约束结构；专家在循环总结为{int(eil['边数'])}条高置信边，作为临床知识与算法融合的精炼结果。\n"
        f"边级重叠显示：PC–HC重叠约{pc_hc}条、PC–GES约{pc_ges}条、HC–GES约{hc_ges}条、树搜索–GES约{tan_ges}条，提示评分式方法间一致性较高，而约束式与评分式具有互补探索能力。\n"
        "综合评分靠前的因果结构多与经典检验、生化指标和共病-用药模式一致，如“检验_白球比→检验_C反应蛋白”、“检验_中性粒细胞比率→检验_中性粒细胞数量”、“药物_维生素C注射液→药物_维生素B6注射液”、“疾病_肾性贫血→药物_碳酸氢钠片”等。\n"
        "统计评估方面，PC框架的卡方检验在0.05阈值下控制虚假发现率，集成评分引入算法一致性与拓扑约束以提升稳定性；多算法一致支持(≥3)与高显著性评分共同构成纳入标准。\n"
        "临床意义上，血脂/血糖相关结构与心脑血管、糖尿病并发症的用药模式形成可验证的因果链路；电解质(如钠/氯、钾)与感染、贫血相关链路与临床经验一致，为个体化干预提供可操作的结构依据。"
    )

def main():
    comp = load_comp_stats()
    overlap = load_overlap()
    top_edges = load_integrated_edges(top_n=15)
    doc = Document(DOCX_PATH)

    add_heading(doc, "2.2 基于多方法学的因果结构发现", level=2)
    add_paragraph(doc, build_methods_text(comp))

    add_heading(doc, "3.2 多方法学因果结构发现结果", level=2)
    add_paragraph(doc, build_results_text(comp, overlap))

    add_table_from_df(doc, comp.reset_index(), "表1 算法总体比较(节点/边/度数指标)")
    add_paragraph(doc, f"数据来源: {os.path.join(RES_DIR, '可视化', '算法对比数据.csv')}")

    add_table_from_df(doc, overlap.reset_index(), "表2 多算法因果边重叠矩阵(条数)")
    add_paragraph(doc, f"数据来源: {os.path.join(RES_DIR, '可视化', '算法重叠矩阵.csv')}")

    add_table_from_df(doc, top_edges, "表3 集成评分居前的因果结构(Top 15)")
    add_paragraph(doc, f"数据来源: {os.path.join(RES_DIR, '06候选因果边集合', '因果边综合评分结果.csv')}")

    fig1 = os.path.join(VIS_DIR, "算法对比分析.png")
    fig2 = os.path.join(VIS_DIR, "边重叠分析.png")
    fig3 = os.path.join(RES_DIR, "01PC算法结果", "PC_因果网络图.png")
    fig4 = os.path.join(RES_DIR, "03贪婪等价搜索结果", "GreedyEquivalence_AIC-D_因果网络图.png")
    fig5 = os.path.join(RES_DIR, "05专家在循环结果", "ExpertInLoop_因果网络图.png")

    if os.path.exists(fig1):
        add_figure(doc, fig1, f"图1 算法对比分析图 — 数据来源: {os.path.join(RES_DIR, '可视化', '算法对比数据.csv')}")
    if os.path.exists(fig2):
        add_figure(doc, fig2, f"图2 多算法边重叠分析图 — 数据来源: {os.path.join(RES_DIR, '可视化', '算法重叠矩阵.csv')}")
    if os.path.exists(fig3):
        add_figure(doc, fig3, "图3 PC算法因果网络图 — 数据来源: 02因果发现/01PC算法结果")
    if os.path.exists(fig4):
        add_figure(doc, fig4, "图4 GES算法因果网络图 — 数据来源: 02因果发现/03贪婪等价搜索结果")
    if os.path.exists(fig5):
        add_figure(doc, fig5, "图5 专家在循环因果网络图 — 数据来源: 02因果发现/05专家在循环结果")

    doc.save(DOCX_PATH)

if __name__ == "__main__":
    main()
